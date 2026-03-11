import os
import re
import uuid
import shutil
import copy
from flask import Flask, render_template, request, jsonify, send_file
import pdfplumber
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

app = Flask(__name__)

# Use /tmp for file storage on cloud servers (Render, etc.)
# Locally, use project subdirectories
if os.environ.get('RENDER') or os.environ.get('PRODUCTION'):
    app.config['UPLOAD_FOLDER'] = '/tmp/pdf_extractor/uploads'
    app.config['OUTPUT_FOLDER'] = '/tmp/pdf_extractor/output'
else:
    app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')
    app.config['OUTPUT_FOLDER'] = os.path.join(os.path.dirname(__file__), 'output')

app.config['TEMPLATE_PATH'] = os.path.join(os.path.dirname(__file__), 'template', 'Doc modelo.docx')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max

# Ensure directories exist
for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
    os.makedirs(folder, exist_ok=True)


def extract_data_from_pdf(pdf_path):
    """Extract prisoner data from the standard PDF format.
    
    The PDF text follows a consistent pattern:
      - Line N: Label (e.g., "Nome", "Pai", "Mãe", "Nascimento ...")
      - Line N+1: Value (possibly with trailing noise after multiple spaces)
      
    For RG: appears in "RG Nome RG Nome" header, value on next line
    For CPF: appears in a line containing CPF numbers or "NÃO ENCONTRADO"
    """
    data = {
        'NOME': '',
        'RG': '',
        'CPF': '',
        'NOME_PAI': '',
        'NOME_MAE': '',
        'DATA_NASCIMENTO': '',
        'filename': os.path.basename(pdf_path)
    }
    
    # UI elements from the PDF to reject
    GARBAGE_VALUES = {'VOLTAR', 'LIMPAR', 'VOLTAR LIMPAR', 'SAIR', '', '-'}
    
    def clean_value(val):
        """Remove trailing noise (RJ codes, pipes, etc.) from extracted values."""
        if not val:
            return ''
        # Split on double spaces — name is before, noise is after
        val = val.split('  ')[0].strip()
        # Remove trailing pipes and RJ codes
        val = re.sub(r'\s*\|.*$', '', val).strip()
        val = re.sub(r'\s+RRJJ.*$', '', val).strip()
        val = re.sub(r'\s+\d{5,}.*$', '', val).strip()
        return val
    
    def is_valid_name(name):
        """Check if a string looks like a valid person name."""
        if not name:
            return False
        clean = clean_value(name)
        if not clean or clean in GARBAGE_VALUES:
            return False
        if len(clean) < 3:
            return False
        # Must be mostly uppercase letters and spaces
        if re.match(r'^[A-ZÁÉÍÓÚÂÊÎÔÛÃÕÇ\s\.]+$', clean) and ' ' in clean:
            return True
        return False
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            page = pdf.pages[0]
            text = page.extract_text() or ''
            lines = text.split('\n')
            
            # ============================================================
            # PRIMARY STRATEGY: Text line parsing
            # Pattern: label line → value on next line
            # ============================================================
            
            for i, line in enumerate(lines):
                stripped = line.strip()
                next_line = lines[i + 1].strip() if i + 1 < len(lines) else ''
                next_line_2 = lines[i + 2].strip() if i + 2 < len(lines) else ''
                
                # --- NOME ---
                # Pattern 1: "Nome" alone on a line → value on next line
                # Pattern 2: "Nome 44338899 || RRJJ..." → value on next line
                # Skip "RG Nome RG Nome" header and "Nome Social"
                if not data['NOME']:
                    is_nome_label = (
                        (stripped == 'Nome' or stripped.startswith('Nome ')) and
                        'RG' not in stripped and 
                        'Social' not in stripped and
                        'Vulgos' not in stripped
                    )
                    if is_nome_label and next_line and is_valid_name(next_line):
                        data['NOME'] = clean_value(next_line)
                
                # --- RG ---
                if not data['RG'] and 'RG' in stripped and 'Outros' not in stripped:
                    rg_match = re.match(r'^(\d{6,12})\b', next_line)
                    if rg_match:
                        data['RG'] = rg_match.group(1)
                
                # --- PAI ---
                # Pattern 1: "Pai" alone → value on next line
                # Pattern 2: "Pai 00111111..." → value on next line
                if not data['NOME_PAI']:
                    if (stripped == 'Pai' or stripped.startswith('Pai ')) and 'Dados' not in stripped:
                        if next_line and is_valid_name(next_line):
                            data['NOME_PAI'] = clean_value(next_line)
                
                # --- MÃE ---
                # Pattern 1: "Mãe" alone → value on next line
                # Pattern 2: "Mãe 338899 || ..." → value on next line
                if not data['NOME_MAE']:
                    is_mae_label = False
                    for mae_variant in ['Mãe', 'Mae', 'Mâe']:
                        if stripped == mae_variant or stripped.startswith(mae_variant + ' '):
                            is_mae_label = True
                            break
                    # Also handle encoding: short label starting with M ending with e
                    if not is_mae_label and len(stripped) >= 2 and len(stripped.split()[0]) <= 4:
                        first_word = stripped.split()[0]
                        if first_word.startswith('M') and first_word.endswith('e') and first_word != 'Nome':
                            is_mae_label = True
                    
                    if is_mae_label and next_line and is_valid_name(next_line):
                        data['NOME_MAE'] = clean_value(next_line)
                
                # --- NASCIMENTO ---
                # Line: "Nascimento Idade aproximada Nacionalidade Naturalidade ..."
                # Some PDFs: date on next line, others: noise line then date 2 lines ahead
                if not data['DATA_NASCIMENTO'] and 'Nascimento' in stripped and 'Data' not in stripped:
                    # Check same line
                    date_match = re.search(r'(\d{2}/\d{2}/\d{4})', stripped)
                    if date_match:
                        data['DATA_NASCIMENTO'] = date_match.group(1)
                    # Check next line
                    if not data['DATA_NASCIMENTO'] and next_line:
                        date_match = re.search(r'(\d{2}/\d{2}/\d{4})', next_line)
                        if date_match:
                            data['DATA_NASCIMENTO'] = date_match.group(1)
                    # Check 2 lines ahead (some PDFs have a noise line in between)
                    if not data['DATA_NASCIMENTO'] and next_line_2:
                        date_match = re.search(r'(\d{2}/\d{2}/\d{4})', next_line_2)
                        if date_match:
                            data['DATA_NASCIMENTO'] = date_match.group(1)
                
                # --- CPF ---
                # Line containing CPF values or "NÃO ENCONTRADO"
                # Format: "CPF NÃO ENCONTRADO 314.576.888-32 314.576.888-32 CPF NÃO ENCONTRADO"
                # Order: DETRAN | SEAP | RECEITA FEDERAL | INTERNO
                if not data['CPF'] and 'CPF' in stripped:
                    cpf_numbers = re.findall(r'\d{3}\.\d{3}\.\d{3}-\d{2}', stripped)
                    if cpf_numbers:
                        # Use the first valid CPF found
                        data['CPF'] = cpf_numbers[0]
                    elif 'ENCONTRADO' in stripped.upper():
                        # Only set NÃO ENCONTRADO if this is the data line (not just the header)
                        # The data line typically has more content than just "CPF"
                        if len(stripped) > 5:
                            data['CPF'] = 'NÃO ENCONTRADO'
            
            # ============================================================
            # FALLBACK: Table-based extraction (for any fields still missing)
            # ============================================================
            tables = page.extract_tables()
            
            for table in tables:
                if len(table) < 10:
                    continue
                
                for ri, row in enumerate(table):
                    if not row or not row[0]:
                        continue
                    cell = str(row[0]).strip()
                    
                    # RG from table
                    if not data['RG'] and re.match(r'^\d{6,12}$', cell) and ri <= 3:
                        data['RG'] = cell
                    
                    # Nome from table: "Nome\nVALUE"
                    if not data['NOME'] and '\n' in cell:
                        parts = cell.split('\n', 1)
                        if parts[0].strip() == 'Nome' and is_valid_name(parts[1]):
                            data['NOME'] = clean_value(parts[1])
                    
                    # Pai from table: "Pai\nVALUE"
                    if not data['NOME_PAI'] and '\n' in cell:
                        parts = cell.split('\n', 1)
                        if parts[0].strip() == 'Pai' and is_valid_name(parts[1]):
                            data['NOME_PAI'] = clean_value(parts[1])
                    
                    # Mãe from table (various encodings)
                    if not data['NOME_MAE'] and '\n' in cell:
                        parts = cell.split('\n', 1)
                        label = parts[0].strip()
                        is_mae = (
                            label == 'Mãe' or label == 'Mae' or label == 'Mâe' or
                            (len(label) <= 4 and label.startswith('M') and label.endswith('e') and label != 'Nome')
                        )
                        if is_mae and is_valid_name(parts[1]):
                            data['NOME_MAE'] = clean_value(parts[1])
                    
                    # Nascimento from table: date in next row
                    if not data['DATA_NASCIMENTO']:
                        date_match = re.match(r'^(\d{2}/\d{2}/\d{4})$', cell)
                        if date_match:
                            data['DATA_NASCIMENTO'] = date_match.group(1)
                    
                    if not data['DATA_NASCIMENTO'] and cell == 'Nascimento':
                        if ri + 1 < len(table) and table[ri + 1] and table[ri + 1][0]:
                            next_cell = str(table[ri + 1][0]).strip()
                            date_match = re.search(r'(\d{2}/\d{2}/\d{4})', next_cell)
                            if date_match:
                                data['DATA_NASCIMENTO'] = date_match.group(1)
            
            # CPF fallback in tables
            if not data['CPF']:
                for table in tables:
                    for row in table:
                        for cell in row:
                            if cell and 'CPF' in str(cell):
                                cpf_numbers = re.findall(r'\d{3}\.\d{3}\.\d{3}-\d{2}', str(cell))
                                if cpf_numbers:
                                    data['CPF'] = cpf_numbers[0]
                                    break
                                elif 'ENCONTRADO' in str(cell).upper() and len(str(cell)) > 10:
                                    data['CPF'] = 'NÃO ENCONTRADO'
                                    break
            
            if not data['CPF']:
                data['CPF'] = 'NÃO ENCONTRADO'
            
            # Final validation: clean garbage values
            for key in ['NOME', 'NOME_PAI', 'NOME_MAE']:
                if data[key] and data[key].strip() in GARBAGE_VALUES:
                    data[key] = ''
                
    except Exception as e:
        data['error'] = str(e)
    
    return data


def replace_markers_in_paragraph(paragraph, data):
    """Replace markers in a paragraph while preserving formatting.
    
    The new template has markers in their own runs with bold formatting,
    so we replace within each run individually to preserve formatting.
    """
    full_text = paragraph.text
    if not any(f'{{{key}}}' in full_text for key in data) and '{NOME_MÃE}' not in full_text:
        return
    
    # Build replacement map
    replacements = {}
    for key, value in data.items():
        if key in ('filename', 'error', '_filepath'):
            continue
        replacements['{' + key + '}'] = str(value)
    # Handle encoding variation for MÃE
    replacements['{NOME_MÃE}'] = data.get('NOME_MAE', '')
    
    # Strategy 1: Try per-run replacement (works when markers are in their own runs)
    for run in paragraph.runs:
        for marker, value in replacements.items():
            if marker in run.text:
                run.text = run.text.replace(marker, value)
    
    # Strategy 2: If markers span across runs, do combined replacement
    # Check if any markers remain after per-run replacement
    remaining_text = ''.join(run.text for run in paragraph.runs)
    has_remaining_markers = any(marker in remaining_text for marker in replacements)
    
    if has_remaining_markers and paragraph.runs:
        # Combined approach: join all runs, replace, put in first run
        combined = ''.join(run.text for run in paragraph.runs)
        for marker, value in replacements.items():
            combined = combined.replace(marker, value)
        paragraph.runs[0].text = combined
        for run in paragraph.runs[1:]:
            run.text = ''


def generate_document(all_data):
    """Generate a DOCX document with all the data from multiple PDFs."""
    template_path = app.config['TEMPLATE_PATH']
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
    
    output_filename = f"documento_gerado_{uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
    
    # Load and modify the template for the first entry
    master_doc = Document(template_path)
    
    # Replace markers in the first copy
    first_data = all_data[0]
    _replace_in_doc(master_doc, first_data)
    
    # For additional entries, append the template content
    for i in range(1, len(all_data)):
        # Add a page break before appending new content
        _append_template_copy(master_doc, template_path, all_data[i])
    
    master_doc.save(output_path)
    return output_filename


def _replace_in_doc(doc, data):
    """Replace all markers in a document."""
    for paragraph in doc.paragraphs:
        replace_markers_in_paragraph(paragraph, data)
    
    # Also replace in headers/footers
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                replace_markers_in_paragraph(paragraph, data)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                replace_markers_in_paragraph(paragraph, data)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_markers_in_paragraph(paragraph, data)


def _append_template_copy(master_doc, template_path, data):
    """Append a copy of the template with replaced markers to the master document.
    
    The template itself already contains empty paragraphs that handle
    page separation, so we just copy all elements directly without
    adding an extra page break.
    """
    # Load a fresh copy of the template
    template_doc = Document(template_path)
    _replace_in_doc(template_doc, data)
    
    # Copy all body elements from template directly
    for element in template_doc.element.body:
        # Skip sectPr (section properties) - we don't want to duplicate those
        if element.tag.endswith('sectPr'):
            continue
        master_doc.element.body.append(copy.deepcopy(element))


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload_pdfs():
    if 'pdfs' not in request.files:
        return jsonify({'error': 'Nenhum arquivo enviado'}), 400
    
    files = request.files.getlist('pdfs')
    if not files:
        return jsonify({'error': 'Nenhum arquivo selecionado'}), 400
    
    results = []
    for file in files:
        if file.filename and file.filename.lower().endswith('.pdf'):
            # Save the uploaded file
            safe_name = f"{uuid.uuid4().hex[:8]}_{file.filename}"
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], safe_name)
            file.save(filepath)
            
            # Extract data
            data = extract_data_from_pdf(filepath)
            data['_filepath'] = safe_name
            results.append(data)
    
    if not results:
        return jsonify({'error': 'Nenhum arquivo PDF válido encontrado'}), 400
    
    return jsonify({'data': results})


@app.route('/generate', methods=['POST'])
def generate():
    try:
        all_data = request.json.get('data', [])
        if not all_data:
            return jsonify({'error': 'Nenhum dado fornecido'}), 400
        
        output_filename = generate_document(all_data)
        return jsonify({'filename': output_filename, 'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/download/<filename>')
def download(filename):
    filepath = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    if not os.path.exists(filepath):
        return jsonify({'error': 'Arquivo não encontrado'}), 404
    return send_file(filepath, as_attachment=True, download_name=filename)


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug = not os.environ.get('RENDER')
    app.run(debug=debug, host='0.0.0.0', port=port)
